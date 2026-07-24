"""
Deep Research Mode — persistent, multi-turn research sessions.

State machine: TOPIC -> DISCUSS -> OUTLINE -> GATHER -> DRAFT -> DONE
State persists to data/research_session.json, so "continue our research"
resumes exactly where you left off — even after a restart.

Finalize produces a real .docx on the Desktop: title page, abstract,
headed sections, inline numbered citations and a References list built
ONLY from actually-fetched sources. It then opens visibly in Word.
"""
import json
import ipaddress
import re
import socket
import time
from pathlib import Path
from urllib.parse import quote_plus, urljoin, urlparse, parse_qs, unquote

import requests
from bs4 import BeautifulSoup, FeatureNotFound

from config import Config
from brain.prompts import (
    RESEARCH_SYSTEM_PROMPT, RESEARCH_OUTLINE_PROMPT, RESEARCH_SECTION_PROMPT,
)

UA = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) JARVIS/1.0"}
STAGES = ["TOPIC", "DISCUSS", "OUTLINE", "GATHER", "DRAFT", "DONE"]


def _soup(markup, parser):
    try:
        return BeautifulSoup(markup, parser)
    except FeatureNotFound:
        return BeautifulSoup(markup, "html.parser")


def clarify_research_topic(topic):
    """Return a search/prompt topic with common ambiguous acronyms expanded.

    The user's wording remains the display title.  Clarification is applied only
    to retrieval and drafting, so a short acronym cannot silently steer source
    selection toward an unrelated domain.
    """
    value = (topic or "").strip()
    if not value:
        return ""
    if re.search(r"\bllms?\b", value, flags=re.I):
        legal_context = re.search(
            r"\b(?:master(?:'s)? of laws?|law degree|legal education|law school|"
            r"jurisprudence|attorney|lawyer)\b", value, flags=re.I,
        )
        if not legal_context:
            value = re.sub(
                r"\bLLM(s?)\b",
                lambda match: (
                    "large language models (LLMs)" if match.group(1)
                    else "large language model (LLM)"
                ),
                value,
                flags=re.I,
            )
    return value


def _working_topic(sess):
    return sess.get("semantic_topic") or clarify_research_topic(sess.get("topic", ""))


def _source_matches_context(sess, hit):
    """Reject obvious homonyms before fetching or citing a search result."""
    display_topic = sess.get("topic", "")
    semantic_topic = _working_topic(sess)
    searchable = " ".join((
        hit.get("title", ""), hit.get("snippet", ""), hit.get("url", ""),
    )).lower()

    # LLM defaults to the technology meaning unless the user supplied explicit
    # legal-degree context.  This catches results such as "Legal education in
    # Hong Kong" that otherwise match only the generic word "education".
    if semantic_topic != display_topic and re.search(r"\bllms?\b", display_topic, re.I):
        technology_markers = (
            "large language model", "language model", "artificial intelligence",
            "generative ai", "machine learning", "ollama", "local inference",
            "transformer", "hugging face",
        )
        if not any(marker in searchable for marker in technology_markers):
            return False
        legal_markers = ("master of laws", "law degree", "legal education", "law school")
        if any(marker in searchable for marker in legal_markers):
            return False

    stopwords = {
        "about", "analysis", "and", "education", "for", "from", "into",
        "local", "models", "overview", "report", "research", "running",
        "the", "their", "this", "with",
    }
    topic_words = {
        word for word in re.findall(r"[a-z0-9]+", semantic_topic.lower())
        if len(word) >= 3 and word not in stopwords
    }
    return not topic_words or bool(topic_words.intersection(
        re.findall(r"[a-z0-9]+", searchable)
    ))


# ===========================================================================
# Session persistence
# ===========================================================================
def _blank_session(topic):
    return {
        "topic": topic,
        "semantic_topic": clarify_research_topic(topic),
        "stage": "DISCUSS",
        "created_at": time.time(),
        "discussion": [],          # list of {"role": user|jarvis, "text": ...}
        "outline": [],             # list of section titles
        "sources": [],             # list of {"title","url","notes"}
        "draft": {},               # section title -> text
        "abstract": "",
    }


def load_session():
    try:
        if Config.RESEARCH_SESSION_FILE.exists():
            data = json.loads(Config.RESEARCH_SESSION_FILE.read_text(encoding="utf-8"))
            if data.get("topic"):
                return data
    except Exception:
        pass
    return None


def save_session(sess):
    try:
        Config.RESEARCH_SESSION_FILE.write_text(
            json.dumps(sess, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def clear_session():
    try:
        Config.RESEARCH_SESSION_FILE.unlink(missing_ok=True)
    except Exception:
        pass


# ===========================================================================
# Web gathering (DuckDuckGo HTML — no API key)
# ===========================================================================
def wikipedia_search(query, limit=6):
    results = []
    try:
        response = requests.get(
            "https://en.wikipedia.org/w/api.php",
            params={
                "action": "query", "list": "search", "srsearch": query,
                "format": "json", "utf8": 1, "srlimit": limit,
            },
            headers=UA, timeout=15,
        )
        response.raise_for_status()
        for item in response.json().get("query", {}).get("search", []):
            title = BeautifulSoup(
                item.get("title", ""), "html.parser"
            ).get_text(" ", strip=True)
            if not title:
                continue
            results.append({
                "title": title,
                "url": "https://en.wikipedia.org/wiki/" + quote_plus(
                    title.replace(" ", "_")
                ),
                "snippet": BeautifulSoup(
                    item.get("snippet", ""), "html.parser"
                ).get_text(" ", strip=True),
            })
    except Exception:
        return []
    return results


def ddg_search(query, limit=6):
    urls = []
    try:
        resp = requests.get(
            "https://html.duckduckgo.com/html/?q=" + quote_plus(query),
            headers=UA, timeout=15)
        soup = _soup(resp.text, "lxml")
        for a in soup.select("a.result__a"):
            href = a.get("href", "")
            title = a.get_text(" ", strip=True)
            url = href
            if "uddg=" in href:
                try:
                    qs = parse_qs(urlparse(href).query)
                    url = unquote(qs.get("uddg", [href])[0])
                except Exception:
                    url = href
            if url.startswith("http") and "duckduckgo.com" not in url:
                urls.append({"title": title or url, "url": url})
            if len(urls) >= limit:
                break
    except Exception:
        pass
    return urls or bing_search(query, limit=limit)


def research_search(query, limit=6):
    combined = wikipedia_search(query, limit=limit)
    combined.extend(ddg_search(query, limit=limit))
    seen = set()
    results = []
    for item in combined:
        url = item.get("url", "").split("#", 1)[0]
        if not url or url in seen:
            continue
        seen.add(url)
        results.append(item)
        if len(results) >= limit:
            break
    return results


def search_web(query, limit=6):
    """Hermes-pilot wrapper for JARVIS's existing public-source search."""
    query = str(query or "").strip()
    if not query:
        return []
    return research_search(query, limit=max(1, min(int(limit), 10)))


def _require_public_source_url(url):
    """Reject local/private destinations before a Hermes-proposed read."""
    value = str(url or "").strip()
    parsed = urlparse(value)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("source URL must be public HTTP or HTTPS")
    if parsed.username or parsed.password:
        raise ValueError("source URL credentials are forbidden")
    try:
        port = parsed.port or (443 if parsed.scheme == "https" else 80)
    except ValueError as exc:
        raise ValueError("source URL port is invalid") from exc
    if port not in {80, 443}:
        raise ValueError("source URL port is not allowed")
    host = parsed.hostname.casefold().rstrip(".")
    if host == "localhost" or host.endswith(".localhost"):
        raise ValueError("local source URLs are forbidden")
    try:
        addresses = {item[4][0] for item in socket.getaddrinfo(host, port)}
    except OSError as exc:
        raise ValueError("source hostname could not be resolved") from exc
    if not addresses or any(not ipaddress.ip_address(address).is_global for address in addresses):
        raise ValueError("private or non-public source address is forbidden")
    return value


def _fetch_public_page_text(url, max_chars):
    """Fetch text while validating every redirect destination."""
    current = _require_public_source_url(url)
    for _redirect in range(6):
        response = requests.get(
            current, headers=UA, timeout=10, allow_redirects=False,
        )
        if response.status_code in {301, 302, 303, 307, 308}:
            location = str(response.headers.get("Location") or "").strip()
            if not location:
                raise ValueError("source redirect has no destination")
            current = _require_public_source_url(urljoin(current, location))
            continue
        response.raise_for_status()
        if "text" not in response.headers.get("Content-Type", "text"):
            return current, ""
        soup = _soup(response.text, "lxml")
        for tag in soup([
            "script", "style", "noscript", "header", "footer", "nav",
            "aside", "form", "iframe",
        ]):
            tag.decompose()
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        text = "\n".join(paragraph for paragraph in paragraphs if len(paragraph) > 40)
        return current, text[:max_chars]
    raise ValueError("source exceeded the redirect limit")


def read_source(url, max_chars=3500):
    """Read bounded public text through the existing research fetcher."""
    bounded = max(200, min(int(max_chars), 6000))
    final_url, text = _fetch_public_page_text(url, bounded)
    return {"url": final_url, "text": text}


def summarize_sources(sources, topic="", max_chars=2400):
    """Create a deterministic evidence digest; never execute returned text."""
    if not isinstance(sources, list):
        raise ValueError("sources must be a list")
    lines = []
    if str(topic or "").strip():
        lines.append(f"Source summary for {str(topic).strip()}:")
    for index, source in enumerate(sources[:10], 1):
        if not isinstance(source, dict):
            continue
        title = re.sub(r"\s+", " ", str(source.get("title") or f"Source {index}")).strip()
        evidence = re.sub(
            r"\s+", " ",
            str(source.get("notes") or source.get("snippet") or source.get("text") or ""),
        ).strip()
        url = str(source.get("url") or "").strip()
        if evidence:
            lines.append(f"[{index}] {title}: {evidence[:500]}" + (f" ({url})" if url else ""))
    return "\n".join(lines)[:max(200, min(int(max_chars), 6000))]


def bing_search(query, limit=6):
    results = []
    try:
        resp = requests.get(
            "https://www.bing.com/search?format=rss&q=" + quote_plus(query),
            headers=UA, timeout=15)
        resp.raise_for_status()
        soup = _soup(resp.content, "xml")
        for item in soup.find_all("item"):
            url = item.link.get_text(strip=True) if item.link else ""
            title = item.title.get_text(" ", strip=True) if item.title else url
            snippet = item.description.get_text(" ", strip=True) if item.description else ""
            if url.startswith("http"):
                results.append({"title": title or url, "url": url,
                                "snippet": snippet})
            if len(results) >= limit:
                break
    except Exception:
        return []
    return results


def fetch_page_text(url, max_chars=3500):
    try:
        resp = requests.get(url, headers=UA, timeout=10, allow_redirects=True)
        if "text" not in resp.headers.get("Content-Type", "text"):
            return ""
        soup = _soup(resp.text, "lxml")
        for tag in soup(["script", "style", "noscript", "header", "footer",
                         "nav", "aside", "form", "iframe"]):
            tag.decompose()
        paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
        text = "\n".join(p for p in paragraphs if len(p) > 40)
        return text[:max_chars]
    except Exception:
        return ""


def gather_sources(sess, ctx, max_sources=10, progress_cb=None, checkpoint=None,
                   summarize_with_llm=True):
    """Search + fetch + take notes per source. Updates sess['sources']."""
    working_topic = _working_topic(sess)
    queries = [working_topic]
    for sec in sess.get("outline", [])[:3]:
        queries.append(f"{working_topic} {sec}")
    if len(queries) < 3:
        queries.append(working_topic + " overview analysis")

    found = []
    seen = set()
    for q in queries:
        if checkpoint:
            checkpoint()
        if progress_cb:
            progress_cb(f"Searching for sources: {q}")
        for hit in research_search(q, limit=5):
            key = hit["url"].split("#")[0]
            if key in seen:
                continue
            seen.add(key)
            found.append(hit)
            if len(found) >= max_sources + 4:
                break
        if len(found) >= max_sources + 4:
            break

    sources = []
    for hit in found:
        if checkpoint:
            checkpoint()
        if len(sources) >= max_sources:
            break
        if progress_cb:
            progress_cb(f"Verifying source {len(sources) + 1}: {hit['title']}")
        if not _source_matches_context(sess, hit):
            continue
        text = fetch_page_text(hit["url"])
        if not text or len(text) < 300:
            text = hit.get("snippet", "")
        if not text or len(text) < 80:
            continue
        notes = ""
        if summarize_with_llm and ctx.llm.available:
            notes = ctx.llm.quick(
                f"Extract 3 short bullet notes from this text, strictly relevant "
                f"to the research topic \"{working_topic}\". Plain text, no markdown.\n\n"
                f"TEXT:\n{text[:2500]}",
                max_tokens=220,
            )
        if not notes:
            notes = text[:400]
        parsed = urlparse(hit["url"])
        sources.append({
            "title": hit["title"],
            "publisher": parsed.netloc.removeprefix("www."),
            "url": hit["url"],
            "publication_date": "Unavailable",
            "access_time": time.strftime("%Y-%m-%dT%H:%M:%S%z"),
            "claim_supported": notes.splitlines()[0][:300] if notes else "",
            "citation_identifier": f"[{len(sources) + 1}]",
            "notes": notes,
        })
    sess["sources"] = sources
    save_session(sess)
    return sources


# ===========================================================================
# Drafting
# ===========================================================================
def _sources_block(sess):
    lines = []
    for i, s in enumerate(sess.get("sources", []), 1):
        lines.append(f"[{i}] {s['title']} — {s['url']}\nNotes: {s['notes']}")
    return "\n\n".join(lines) if lines else "(no sources fetched yet)"


def draft_section(sess, section, ctx, extra=""):
    if not ctx.llm.available:
        sources = sess.get("sources", [])
        if not sources:
            return ""
        paragraphs = [
            f"This section examines {section.lower()} for {_working_topic(sess)} "
            f"using only the verified public sources collected by JARVIS."
        ]
        for index, source in enumerate(sources, 1):
            notes = re.sub(r"\s+", " ", str(source.get("notes") or "")).strip()
            if not notes:
                continue
            excerpt = notes[:420].rsplit(" ", 1)[0] if len(notes) > 420 else notes
            paragraphs.append(f"{excerpt} [{index}]")
        return "\n\n".join(paragraphs)
    prompt = RESEARCH_SECTION_PROMPT.format(
        section=section, topic=_working_topic(sess), extra=extra,
        sources=_sources_block(sess),
    )
    text = ctx.llm.quick(prompt, system=RESEARCH_SYSTEM_PROMPT.format(
        sources=_sources_block(sess)), max_tokens=1500)
    return text.strip()


def draft_all(sess, ctx, progress_cb=None, checkpoint=None, section_cb=None):
    for i, section in enumerate(sess.get("outline", []), 1):
        if checkpoint:
            checkpoint()
        if progress_cb:
            progress_cb(f"Drafting section {i} of {len(sess['outline'])}: {section}.")
        sess["draft"][section] = draft_section(sess, section, ctx)
        save_session(sess)
        if section_cb:
            section_cb(section, sess["draft"][section])
    joined = "\n\n".join(
        f"## {k}\n{v}" for k, v in sess["draft"].items())[:4000]
    if checkpoint:
        checkpoint()
    if ctx.llm.available:
        abstract = ctx.llm.quick(
            f"Write a 120-180 word abstract for this research document titled "
            f"\"{_working_topic(sess)}\". Plain prose, no markdown.\n\n{joined}",
            max_tokens=350)
    else:
        abstract = (
            f"This source-grounded briefing reviews {_working_topic(sess)} "
            f"across {len(sess.get('outline', []))} sections using "
            f"{len(sess.get('sources', []))} verified public sources. "
            "Because no generation provider was configured, JARVIS preserved "
            "retrieved source evidence directly and did not add unsupported claims."
        )
    sess["abstract"] = (abstract or "").strip()
    save_session(sess)


def build_research_session(topic, ctx, max_sources=8, max_sections=6,
                           progress_cb=None, checkpoint=None,
                           summarize_with_llm=True, min_sources=3,
                           section_cb=None):
    """Build a source-grounded in-memory report for visible Office insertion."""
    topic = (topic or "").strip()
    if not topic:
        return None
    sess = _blank_session(topic)
    save_session(sess)
    if progress_cb:
        progress_cb("Creating the report outline")
    _propose_outline(sess, ctx)
    sess["outline"] = sess.get("outline", [])[:max_sections]
    save_session(sess)
    sources = gather_sources(
        sess, ctx, max_sources=max_sources,
        progress_cb=progress_cb, checkpoint=checkpoint,
        summarize_with_llm=summarize_with_llm,
    )
    if len(sources) < min(max_sources, max(1, min_sources)):
        return None
    sess["stage"] = "DRAFT"
    save_session(sess)
    draft_all(
        sess, ctx, progress_cb=progress_cb, checkpoint=checkpoint,
        section_cb=section_cb,
    )
    if not any(text.strip() for text in sess.get("draft", {}).values()):
        return None
    sess["stage"] = "DONE"
    save_session(sess)
    return sess


# ===========================================================================
# Finalize -> .docx on the Desktop, opened in Word
# ===========================================================================
def finalize(sess, ctx, open_visible=True):
    if not sess.get("draft"):
        return None
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(sess["topic"])
    run.bold = True
    run.font.size = Pt(26)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("A collaborative research document").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(time.strftime("%d %B %Y"))
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.add_run("Prepared with JARVIS Research Assistant").italic = True
    doc.add_page_break()

    # ---- abstract
    if sess.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(sess["abstract"])
        doc.add_page_break()

    # ---- sections
    for section, text in sess["draft"].items():
        doc.add_heading(section, level=1)
        for para in text.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    # ---- references (fetched sources ONLY — never invented)
    doc.add_page_break()
    doc.add_heading("References", level=1)
    for i, s in enumerate(sess.get("sources", []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run(f"[{i}] ").bold = True
        p.add_run(f"{s['title']}. ")
        p.add_run(s["url"]).italic = True

    slug = re.sub(r"[^\w\s-]", "", sess["topic"])[:50].strip().replace(" ", "_")
    path = Config.DESKTOP_PATH / f"research_{slug or 'document'}.docx"
    doc.save(str(path))
    if open_visible:
        _open_in_word(str(path), ctx)
    return path


def _open_in_word(path, ctx):
    try:
        import win32com.client as win32
        word = win32.DispatchEx("Word.Application")
        word.Visible = True
        word.Documents.Open(path)

        def _closer(w=word):
            try:
                w.Quit()
            except Exception:
                pass
        ctx.registry.register("document", Path(path).name,
                              window_title=Path(path).stem, closer=_closer,
                              extra={"path": path})
        return True
    except Exception:
        try:
            import os
            os.startfile(path)
            ctx.registry.register("document", Path(path).name,
                                  window_title=Path(path).stem,
                                  extra={"path": path})
        except Exception:
            pass
        return False


# ===========================================================================
# Conversational state machine (driven from main.py while pending)
# ===========================================================================
def _say(sess, role, text):
    sess.setdefault("discussion", []).append({"role": role, "text": text})
    if len(sess["discussion"]) > 40:
        sess["discussion"] = sess["discussion"][-40:]
    save_session(sess)


def start(topic, ctx):
    topic = (topic or "").strip()
    if not topic:
        return "Research what, sir?"
    sess = _blank_session(topic)
    save_session(sess)
    ctx.state["research"] = True
    ctx.pending = {"kind": "research"}

    opening = ctx.llm.quick(
        f"We are starting a research project on \"{topic}\". In under 80 words: "
        f"show you understand the topic, propose two or three sharp angles, "
        f"then ask the user one focused question about their goals or scope.",
        system=RESEARCH_SYSTEM_PROMPT.format(sources="(none yet)"),
        max_tokens=220) if ctx.llm.available else ""
    if not opening:
        opening = (f"Splendid, sir — researching {topic}. Tell me your goals: "
                   f"breadth or depth, academic or practical?")
    _say(sess, "jarvis", opening)
    return f"Let's research {topic}, sir. {opening}"


def resume(ctx):
    sess = load_session()
    if not sess:
        return "We have no research in progress, sir. Say 'research X' to begin one."
    ctx.state["research"] = True
    ctx.pending = {"kind": "research"}
    stage = sess.get("stage", "DISCUSS")
    bits = [f"Resuming our research on {sess['topic']}, sir."]
    bits.append(f"We are at the {stage.lower()} stage.")
    if sess.get("outline"):
        bits.append(f"The outline has {len(sess['outline'])} sections.")
    if sess.get("sources"):
        bits.append(f"We hold {len(sess['sources'])} sources.")
    bits.append("How shall we proceed?")
    return " ".join(bits)


def handle_utterance(text, ctx):
    """Route one utterance inside an active research session."""
    sess = load_session()
    if sess is None:
        ctx.pending = None
        ctx.state.pop("research", None)
        return "The research session seems to have vanished, sir."
    t = (text or "").strip()
    low = t.lower()

    # ---- global session commands -----------------------------------------
    if re.search(r"\b(exit|leave|quit|cancel) (?:the )?research\b", low):
        ctx.pending = None
        ctx.state.pop("research", None)
        return ("Leaving research mode, sir. The session is saved — say "
                "'continue our research' anytime.")
    if re.search(r"\bfinali[sz]e\b", low):
        return _do_finalize(sess, ctx)
    if "what sources" in low or "list sources" in low or "read the sources" in low:
        if not sess.get("sources"):
            return "We haven't gathered sources yet, sir."
        lines = [f"We hold {len(sess['sources'])} sources, sir."]
        for i, s in enumerate(sess["sources"], 1):
            lines.append(f"{i}: {s['title']}.")
        return " ".join(lines)
    if "read me the outline" in low or "read the outline" in low or low == "outline":
        if not sess.get("outline"):
            return "There's no outline yet, sir."
        lines = ["The outline, sir:"]
        for i, s in enumerate(sess["outline"], 1):
            lines.append(f"{i}: {s}.")
        return " ".join(lines)
    m = re.search(r"change (?:the )?focus to (.+)", low)
    if m:
        sess["topic"] = m.group(1).strip()
        save_session(sess)
        return f"Refocused on {sess['topic']}, sir. Shall we revise the outline?"

    stage = sess.get("stage", "DISCUSS")

    # ---- DISCUSS ----------------------------------------------------------
    if stage == "DISCUSS":
        if re.search(r"\b(make|draft|propose|write|give me) (?:an |the )?outline\b", low) \
                or low in ("outline", "outline please", "continue", "go on"):
            return _propose_outline(sess, ctx)
        _say(sess, "user", t)
        reply = ctx.llm.quick(
            f"Topic: {sess['topic']}\nDiscussion so far:\n" +
            "\n".join(f"{d['role']}: {d['text']}" for d in sess["discussion"][-10:]) +
            "\n\nRespond as JARVIS the research partner (under 100 words, spoken style). "
            "Engage with their point, challenge it if weak, and steer toward an outline.",
            system=RESEARCH_SYSTEM_PROMPT.format(sources="(not gathered yet)"),
            max_tokens=260) if ctx.llm.available else ""
        if not reply:
            reply = "Noted, sir. Shall I draft the outline now?"
        _say(sess, "jarvis", reply)
        return reply

    # ---- OUTLINE -----------------------------------------------------------
    if stage == "OUTLINE":
        if re.search(r"\b(looks good|good|perfect|continue|proceed|gather|go ahead|fine|ok|okay)\b", low):
            return _do_gather(sess, ctx)
        m = re.search(r"drop section (\w+)", low)
        if m:
            idx = _word_to_index(m.group(1))
            if idx is not None and 0 <= idx < len(sess["outline"]):
                removed = sess["outline"].pop(idx)
                save_session(sess)
                return (f"Section {m.group(1)} ({removed}) dropped, sir. "
                        f"Anything else, or shall I gather sources?")
        m = re.search(r"add (?:a )?section (?:on|about) (.+)", low)
        if m:
            sess["outline"].append(m.group(1).strip().rstrip("."))
            save_session(sess)
            return (f"Added a section on {m.group(1).strip()}, sir. "
                    f"More edits, or shall I gather sources?")
        revised = ctx.llm.quick(
            f"Current outline for \"{sess['topic']}\":\n" +
            "\n".join(f"{i+1}. {s}" for i, s in enumerate(sess["outline"])) +
            f"\n\nUser request: {t}\n"
            "Apply the request and return the FULL revised outline, numbered "
            "'1. Title — one-line description'. Nothing else.",
            max_tokens=500) if ctx.llm.available else ""
        sections = _parse_outline(revised)
        if sections:
            sess["outline"] = sections
            save_session(sess)
            spoken = " ".join(f"{i+1}: {s}." for i, s in enumerate(sections))
            return f"Revised outline, sir. {spoken} Shall I gather sources?"
        return ("I didn't catch an outline edit, sir. Try 'drop section 2', "
                "'add a section on X', or say 'gather sources'.")

    # ---- GATHER ------------------------------------------------------------
    if stage == "GATHER":
        if re.search(r"\b(gather|search|find sources|go|continue|proceed)\b", low):
            return _do_gather(sess, ctx)
        return "Say 'gather sources' and I'll start collecting, sir."

    # ---- DRAFT --------------------------------------------------------------
    if stage == "DRAFT":
        if re.search(r"\b(write|draft|start writing|continue|proceed)\b", low):
            return _do_draft(sess, ctx)
        m = re.search(r"(?:expand|rewrite|redo) (?:section )?(\w+)?", low)
        if m and sess.get("draft"):
            target = m.group(1)
            section = None
            if target:
                idx = _word_to_index(target)
                if idx is not None and 0 <= idx < len(sess["outline"]):
                    section = sess["outline"][idx]
            if section is None:
                section = list(sess["draft"].keys())[-1]
            sess["draft"][section] = draft_section(
                sess, section, ctx, extra=f"Rewrite instruction: {t}")
            save_session(sess)
            return f"Section '{section}' rewritten, sir. More edits, or 'finalize'?"
        if "more academic" in low or "add statistics" in low or "make it" in low:
            if sess.get("draft"):
                section = list(sess["draft"].keys())[-1]
                sess["draft"][section] = draft_section(
                    sess, section, ctx, extra=f"Style instruction: {t}")
                save_session(sess)
                return "Adjusted the latest section as requested, sir."
        return ("We're at the drafting stage, sir. Say 'write the draft', "
                "give me an edit like 'expand section 2', or say 'finalize'.")

    # ---- DONE -----------------------------------------------------------------
    if stage == "DONE":
        return ("This research is finalized, sir. The document is on your desktop. "
                "Say 'exit research' or start a new topic anytime.")

    return "I'm not sure how that fits our research, sir."


# ===========================================================================
# Stage transitions
# ===========================================================================
def _propose_outline(sess, ctx):
    discussion = "\n".join(
        f"{d['role']}: {d['text']}" for d in sess.get("discussion", [])[-12:])
    text = ctx.llm.quick(
        RESEARCH_OUTLINE_PROMPT.format(topic=_working_topic(sess), discussion=discussion),
        system=RESEARCH_SYSTEM_PROMPT.format(sources="(not gathered yet)"),
        max_tokens=600) if ctx.llm.available else ""
    sections = _parse_outline(text)
    if not sections:
        sections = [f"Background and context of {sess['topic']}",
                    "Key developments and current state",
                    "Major debates and open questions",
                    "Analysis and implications",
                    "Conclusions and recommendations"]
    sess["outline"] = sections
    sess["stage"] = "OUTLINE"
    save_session(sess)
    spoken = " ".join(f"{i+1}: {s}." for i, s in enumerate(sections))
    return (f"Here's my proposed outline, sir. {spoken} "
            f"Say things like 'drop section 2' or 'add a section on X', "
            f"or 'gather sources' when you're happy.")


def _do_gather(sess, ctx):
    sess["stage"] = "GATHER"
    save_session(sess)
    ctx.speaker.speak("Gathering sources now, sir. Give me a moment.")
    sources = gather_sources(sess, ctx, max_sources=10)
    if not sources:
        sess["stage"] = "OUTLINE"
        save_session(sess)
        return ("The web search came back empty, sir — possibly no connection. "
                "We can try again or adjust the topic.")
    sess["stage"] = "DRAFT"
    save_session(sess)
    names = "; ".join(s["title"] for s in sources[:4])
    return (f"I've collected {len(sources)} sources, sir — including {names}. "
            f"Say 'write the draft' and I'll compose the document section by section.")


def _do_draft(sess, ctx):
    if not sess.get("sources"):
        return _do_gather(sess, ctx)
    ctx.speaker.speak("Drafting the document, sir. I'll narrate as I go.")
    draft_all(sess, ctx, progress_cb=lambda msg: ctx.speaker.speak(msg))
    sess["stage"] = "DRAFT"
    save_session(sess)
    return ("The full draft is written, sir. Ask me to expand or restyle any "
            "section, or say 'finalize the research' for the Word document.")


def _do_finalize(sess, ctx):
    if not sess.get("draft"):
        if sess.get("sources"):
            return _do_draft(sess, ctx)
        return "There's nothing to finalize yet, sir — no draft exists."
    ctx.speaker.speak("Finalizing the document now, sir.")
    path = finalize(sess, ctx)
    if path is None:
        return "I couldn't build the document, sir."
    sess["stage"] = "DONE"
    save_session(sess)
    n_words = sum(len(v.split()) for v in sess["draft"].values())
    return (f"Done, sir. The document is on your desktop as {Path(path).name} "
            f"and open in Word. It runs about {n_words} words across "
            f"{len(sess['draft'])} sections, citing {len(sess['sources'])} real "
            f"sources. A pleasure researching with you.")


def create_report(topic, ctx):
    topic = (topic or "").strip()
    if not topic:
        return "Research what, sir?"
    sess = _blank_session(topic)
    save_session(sess)
    _propose_outline(sess, ctx)
    sources = gather_sources(sess, ctx, max_sources=10)
    if not sources:
        return f"I couldn't find reliable sources for {topic}, sir."
    draft_all(sess, ctx)
    path = finalize(sess, ctx)
    if path is None or not Path(path).exists():
        return f"I couldn't create the research report about {topic}, sir."
    sess["stage"] = "DONE"
    save_session(sess)
    return f"Created {Path(path).name} with references and opened it in Microsoft Word, sir."


def prepare_report(topic, ctx):
    topic = (topic or "").strip()
    if not topic:
        return "Research what, sir?"
    sess = _blank_session(topic)
    save_session(sess)
    _propose_outline(sess, ctx)
    return f"Prepared a structured outline for {topic}, sir."


def gather_report(ctx):
    sess = load_session()
    if not sess:
        return "No research report is prepared, sir."
    sources = gather_sources(sess, ctx, max_sources=10)
    if not sources:
        return f"I couldn't find reliable sources for {sess['topic']}, sir."
    return f"Gathered {len(sources)} real sources for {sess['topic']}, sir."


def draft_report(ctx):
    sess = load_session()
    if not sess or not sess.get("sources"):
        return "The report has no sources to draft from, sir."
    draft_all(sess, ctx)
    return f"Drafted the structured report about {sess['topic']}, sir."


def finalize_report(ctx):
    sess = load_session()
    if not sess or not sess.get("draft"):
        return "The report has no completed draft, sir."
    path = finalize(sess, ctx, open_visible=False)
    if path is None or not Path(path).exists():
        return "I couldn't save the research report, sir."
    ctx.state["last_research_path"] = str(path)
    sess["stage"] = "DONE"
    save_session(sess)
    return f"Saved {Path(path).name} with references, sir."


def open_report(ctx):
    path = Path(ctx.state.get("last_research_path", ""))
    if not path.is_file():
        return "I can't find the completed research report, sir."
    if _open_in_word(str(path), ctx):
        return f"Opened {path.name} visibly in Microsoft Word, sir."
    return f"Windows could not open {path.name} in Microsoft Word, sir."


# ===========================================================================
# Small helpers
# ===========================================================================
def _parse_outline(text):
    if not text:
        return []
    sections = []
    for line in text.splitlines():
        line = line.strip()
        m = re.match(r"^\d+[\.\)]\s*(.+)$", line)
        if m:
            sections.append(m.group(1).strip())
    return sections[:10]


def _word_to_index(word):
    mapping = {"one": 0, "first": 0, "1": 0, "two": 1, "second": 1, "2": 1,
               "three": 2, "third": 2, "3": 2, "four": 3, "fourth": 3, "4": 3,
               "five": 4, "fifth": 4, "5": 4, "six": 5, "sixth": 5, "6": 5,
               "seven": 6, "seventh": 6, "7": 6, "eight": 7, "eighth": 7, "8": 7,
               "nine": 8, "ninth": 8, "9": 8, "ten": 9, "tenth": 9, "10": 9}
    return mapping.get(str(word).strip().lower())


# ===========================================================================
# Skill dispatch entry
# ===========================================================================
def handle(intent, ctx):
    skill = intent.get("skill")
    params = intent.get("params", {}) or {}
    if skill == "research.start":
        return start(params.get("topic", ""), ctx)
    if skill == "research.create_report":
        return create_report(params.get("topic", ""), ctx)
    if skill == "research.prepare_report":
        return prepare_report(params.get("topic", ""), ctx)
    if skill == "research.gather_report":
        return gather_report(ctx)
    if skill == "research.draft_report":
        return draft_report(ctx)
    if skill == "research.finalize_report":
        return finalize_report(ctx)
    if skill == "research.open_report":
        return open_report(ctx)
    if skill == "research.continue":
        return resume(ctx)
    if skill == "research.finalize":
        sess = load_session()
        if sess is None:
            return "No research session exists, sir."
        return _do_finalize(sess, ctx)
    if skill == "research.outline":
        sess = load_session()
        if sess is None:
            return "No research session exists, sir."
        return _propose_outline(sess, ctx)
    return None
