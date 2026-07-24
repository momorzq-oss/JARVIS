"""
Microsoft Word — write documents on the Desktop with python-docx, then
open them VISIBLY in Word via win32com so you watch them appear.
"continue writing" appends AI-generated content to an existing file.
Every opened document is registered so "close Word" works.
"""
import re
import time
import uuid
from difflib import get_close_matches
from pathlib import Path

from config import Config
from brain.prompts import WORD_WRITE_PROMPT
from core.live_task import LIVE_INTERACTIVE, TaskCancelled
from core.save_workflow import PendingSaveRequest


# ---------------------------------------------------------------------------
# Word COM helpers
# ---------------------------------------------------------------------------
def _open_in_word(path, ctx):
    try:
        import win32com.client as win32
        word = win32.DispatchEx("Word.Application")
        word.Visible = True
        word.Documents.Open(str(path))

        def _closer(w=word):
            try:
                w.Quit()
            except Exception:
                pass
        ctx.registry.register("document", Path(path).name,
                              window_title=Path(path).stem, closer=_closer,
                              extra={"path": str(path)})
        return True
    except Exception:
        try:
            import os
            os.startfile(str(path))
            ctx.registry.register("document", Path(path).name,
                                  window_title=Path(path).stem,
                                  extra={"path": str(path)})
        except Exception:
            pass
        return False


def _find_docx(name):
    name_l = (name or "").lower().replace(".docx", "").replace(".doc", "")
    candidates = []
    try:
        for p in Config.DESKTOP_PATH.glob("*.doc*"):
            if p.suffix.lower() in (".docx", ".doc"):
                candidates.append(p)
    except Exception:
        pass
    if not candidates:
        return None
    for p in candidates:
        if name_l in p.stem.lower():
            return p
    matches = get_close_matches(name_l, [p.stem.lower() for p in candidates],
                                n=1, cutoff=0.4)
    if matches:
        for p in candidates:
            if p.stem.lower() == matches[0]:
                return p
    return None


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------
def _slug(text):
    s = re.sub(r"[^\w\s-]", "", text)[:60].strip().replace(" ", "_")
    return s or "document"


def _build_docx(markdownish, path):
    """Parse '# Title' / '## Heading' / paragraphs into a styled .docx."""
    from docx import Document
    doc = Document()
    for raw in markdownish.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
        else:
            doc.add_paragraph(stripped)
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Actions
# ---------------------------------------------------------------------------
def create_document(ctx):
    """Open Word visibly with a new unsaved document."""
    try:
        from skills.office_service import WordService
        service = WordService()
        service.open(visible=True)
        service.new_document()

        def _closer(svc=service):
            svc.close(save=False)

        entry = ctx.registry.register(
            "document", "New Word Document", window_title="Word",
            pid=service.process_id, hwnd=service.window_handle,
            closer=_closer, extra={
                "unsaved": "true", "owner": "jarvis",
                "application": "Microsoft Word", "process_name": "WINWORD.EXE",
            },
        )
        ctx.state["word_service"] = service
        ctx.state["active_office_entry"] = entry["id"]
        return "A new Word document is open, sir."
    except Exception as exc:
        return f"I couldn't create a Word document: {exc}."


def insert_text(text, ctx):
    """Insert user-supplied text locally into the active JARVIS Word document."""
    value = str(text or "").strip()
    if not value:
        return "What text should I insert, sir?"
    service = ctx.state.get("word_service")
    if service is None:
        return "No JARVIS Word document is active, sir."
    try:
        service.insert_text(value)
        entry_id = ctx.state.get("active_office_entry")
        if entry_id:
            ctx.registry.update_entry(entry_id, unsaved_state="unsaved")
        return f"Inserted {len(value.split())} words into Word, sir."
    except Exception as exc:
        return f"Word text insertion failed: {exc}."


def save_document(path, ctx):
    """Save the active JARVIS Word document to an explicit local path."""
    service = ctx.state.get("word_service")
    if service is None:
        return "No JARVIS Word document is active, sir."
    target = Path(str(path or "")).expanduser()
    if not target.is_absolute():
        target = Config.SOURCE_DIR / target
    if target.suffix.lower() not in {".docx", ".doc"}:
        target = target.with_suffix(".docx")
    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        service.save(target)
        if not target.is_file():
            return "Word reported a save, but the file was not verified, sir."
        entry_id = ctx.state.get("active_office_entry")
        if entry_id:
            ctx.registry.update_entry(
                entry_id, name=target.name, display_name=target.name,
                window_title=target.stem, file_path=str(target),
                unsaved_state="saved",
            )
        ctx.state["last_generated_output"] = str(target)
        pending = getattr(ctx, "pending", None)
        if isinstance(pending, dict) and pending.get("kind") == "save_document":
            ctx.pending = None
        return f"Saved and verified at {target}, sir."
    except Exception as exc:
        return f"The Word document was not saved: {exc}."


def create_live_document(topic, ctx, report_length="full"):
    """Draft and insert a report progressively through visible Word COM."""
    topic = (topic or "").strip()
    if not topic:
        return "What should the report be about, sir?"
    task = getattr(ctx, "live_task", None)
    task_id = uuid.uuid4().hex[:12]
    if task is not None:
        task.start(task_id, f"Live Word report: {topic}",
                   application="Microsoft Word", mode=LIVE_INTERACTIVE)
        task.state.delay_ms = max(0, Config.LIVE_TYPING_DELAY_MS)

    try:
        short_report = str(report_length).lower() == "short"
        speaker = getattr(ctx, "speaker", None)
        if speaker is not None:
            speaker.speak(f"Opening Word and researching {topic}, sir.")

        # A live request must become visible before network research begins.
        # Previously Word was opened only after every source and section had
        # been prepared, leaving the desktop unchanged for over a minute.
        from skills.office_service import WordService
        service = WordService()
        service.open(visible=True)
        document = service.new_document()
        entry = ctx.registry.register(
            "document", f"{topic.title()} Report", window_title="Word",
            pid=service.process_id, hwnd=service.window_handle,
            closer=lambda svc=service: svc.close(save=False),
            extra={
                "unsaved": "true", "owner": "jarvis",
                "application": "Microsoft Word", "process_name": "WINWORD.EXE",
            },
        )
        ctx.state["word_service"] = service
        ctx.state["active_office_entry"] = entry["id"]
        service.insert_heading(topic.title(), level=1, doc=document)
        words_written = len(topic.split())
        paragraphs_written = 0
        if task is not None:
            task.update(
                step="Microsoft Word opened; researching verified sources",
                progress=3, current_heading=topic.title(),
                words_written=words_written, paragraph_count=0,
            )

        from skills import research

        def research_progress(step):
            if task is not None:
                task.checkpoint()
                current = task.snapshot().get("progress", 0)
                task.update(step=step, progress=min(45, max(3, current + 3)))

        def write_section(section, text):
            """Append each completed draft section immediately in visible Word."""
            nonlocal words_written, paragraphs_written
            if task is not None:
                task.checkpoint()
            service.insert_heading(section, level=1, doc=document)
            words_written += len(section.split())
            for paragraph in re.split(r"\n\s*\n", str(text or "")):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                if task is not None:
                    task.checkpoint()
                service.type_visibly(paragraph, doc=document)
                words_written += len(paragraph.split())
                paragraphs_written += 1
                if task is not None and task.state.delay_ms:
                    time.sleep(task.state.delay_ms / 1000.0)
            if task is not None:
                task.update(
                    step=f"Wrote {section}", progress=min(88, 52 + paragraphs_written * 4),
                    current_heading=section, words_written=words_written,
                    paragraph_count=paragraphs_written,
                )

        session = research.build_research_session(
            topic, ctx,
            max_sources=4 if short_report else 6,
            max_sections=3 if short_report else 5,
            progress_cb=research_progress,
            checkpoint=task.checkpoint if task is not None else None,
            # Raw verified excerpts already ground the section prompts. Avoid
            # one extra cloud request per source before visible writing starts.
            summarize_with_llm=False,
            section_cb=write_section,
        )
        if session is None:
            detail = getattr(ctx.llm, "last_error", "") or "No verified sources or draft"
            if task is not None:
                task.fail(detail)
            return "The research service could not build a source-grounded report, sir."
        sources = session.get("sources", [])
        if task is not None:
            task.update(
                step=f"Finishing the report with {len(sources)} verified sources",
                progress=90, sources_found=len(sources), sources_verified=len(sources),
            )
        operations = []
        if session.get("abstract"):
            operations.extend((
                ("heading", "Summary"),
                ("paragraph", session["abstract"]),
            ))
        operations.append(("heading", "References"))
        for source in sources:
            operations.append((
                "paragraph",
                f"{source['citation_identifier']} {source['title']}. "
                f"{source['publisher']}. {source['url']} (accessed {source['access_time']}).",
            ))

        total = max(1, len(operations))
        for index, (kind, text) in enumerate(operations, 1):
            if task is not None:
                task.checkpoint()
                task.update(
                    step=text[:80], progress=90 + int(index * 9 / total),
                    current_heading=text[:120] if kind in ("title", "heading") else "",
                    words_written=words_written, paragraph_count=paragraphs_written,
                )
            if kind in ("title", "heading"):
                service.insert_heading(text, level=1, doc=document)
            else:
                service.type_visibly(text, doc=document)
                paragraphs_written += 1
            words_written += len(text.split())
            if task is not None and task.state.delay_ms:
                time.sleep(task.state.delay_ms / 1000.0)

        def _save(path, svc=service, doc=document, entry_id=entry["id"]):
            svc.save(path, doc=doc)
            ctx.registry.update_entry(
                entry_id, name=Path(path).name, display_name=Path(path).name,
                window_title=Path(path).stem, file_path=str(path),
                unsaved_state="saved",
            )

        request = PendingSaveRequest(
            task_id=task_id,
            document_type="Word",
            suggested_filename=f"{topic.title()} Report",
            suggested_extension=".docx",
            current_application="Microsoft Word",
            save_callback=_save,
        )
        ctx.pending = {"kind": "save_document", "request": request}
        if task is not None:
            task.complete()
        return "The document is ready. Where would you like me to save it?"
    except TaskCancelled:
        return "The live Word task was cancelled, sir."
    except Exception as exc:
        if task is not None:
            task.fail(exc)
        return f"The live Word task failed: {exc}."


def write_document(topic, extra, ctx):
    topic = (topic or "").strip()
    if not topic:
        return "Write about what, sir?"
    if not ctx.llm.available:
        return "My writing brain needs the OpenRouter key, sir."

    ctx.speaker.speak(f"Writing about {topic}, sir. One moment.")
    text = ctx.llm.quick(
        WORD_WRITE_PROMPT.format(topic=topic, extra=extra or ""),
        max_tokens=3000)
    if not text:
        return "The words escaped me, sir — the drafting service failed."
    if not text.lstrip().startswith("#"):
        text = f"# {topic.title()}\n\n" + text

    path = Config.DESKTOP_PATH / f"{_slug(topic)}.docx"
    try:
        _build_docx(text, path)
    except Exception as exc:
        return f"I couldn't build the document: {exc}."

    _open_in_word(path, ctx)
    words = len(text.split())
    return (f"Done, sir. {path.name} is on your desktop and open in Word — "
            f"about {words} words.")


def continue_document(file_name, instruction, ctx):
    path = _find_docx(file_name)
    if path is None:
        return f"I can't find a Word file called {file_name} on your desktop, sir."
    if not ctx.llm.available:
        return "My writing brain needs the OpenRouter key, sir."

    try:
        from docx import Document
        doc = Document(str(path))
        tail = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[-1200:]
    except Exception as exc:
        return f"I couldn't open {path.name}: {exc}."

    prompt = (f"Continue this document naturally in the same style. "
              f"{instruction or ''}\n\nDocument so far (tail):\n{tail}\n\n"
              "Write the continuation only — use '## Heading' for any new "
              "sections, plain paragraphs otherwise. No repetition of the tail.")
    addition = ctx.llm.quick(prompt, max_tokens=1500)
    if not addition:
        return "The drafting service failed, sir."

    try:
        for raw in addition.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=1)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip(), level=1)
            else:
                doc.add_paragraph(stripped)
        doc.save(str(path))
    except Exception as exc:
        return f"I couldn't append to {path.name}: {exc}."

    _open_in_word(path, ctx)
    return f"Added to {path.name} and reopened it, sir."


# ---------------------------------------------------------------------------
# Skill dispatch entry
# ---------------------------------------------------------------------------
def handle(intent, ctx):
    skill = intent.get("skill")
    params = intent.get("params", {}) or {}
    if skill == "office_word.create_document":
        return create_document(ctx)
    if skill == "office_word.create_research_document":
        return create_live_document(
            params.get("topic", ""), ctx,
            report_length=params.get("report_length", "full"),
        )
    if skill == "office_word.insert_text":
        return insert_text(params.get("text", ""), ctx)
    if skill == "office_word.save_document":
        return save_document(params.get("path", ""), ctx)
    if skill == "word.write":
        return write_document(params.get("topic", ""), params.get("extra", ""), ctx)
    if skill == "word.continue":
        return continue_document(params.get("file", ""),
                                 params.get("instruction", ""), ctx)
    return None
